import csv
from io import StringIO
from fpdf import FPDF
import fitz
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_to_csv(response):
    rows = [row.split(',') for row in response.split('\n')]
    csv_data = StringIO()
    csv_writer = csv.writer(csv_data)
    csv_writer.writerows(rows)
    csv_content = csv_data.getvalue()
    return csv_content

def decode_pdf(file_path):
    with fitz.open(file_path) as pdf_document:
        pdf_text = ''
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            pdf_text += page.get_text()
        return pdf_text

def decode_docx(file_path):
    doc = Document(file_path)
    doc_text = ''
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + '\n'
    return doc_text

def decode_default(file_path):
    with open(file_path, "rb") as file_content:
        return file_content.read().decode('utf-8', 'ignore')

def convert_to_csv(response):
    rows = [row.split(',') for row in response.split('\n')]
    csv_data = StringIO()
    csv_writer = csv.writer(csv_data)
    csv_writer.writerows(rows)
    return csv_data.getvalue()

def generate_csv(response, csv_filename):
    csv_content = convert_to_csv(response)
    with open(csv_filename, mode="w", newline="", encoding="utf-8") as csv_file:
        csv_file.write(csv_content)

def generate_pdf(messages, pdf_filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", style="B")

    for msg in messages[1:]:
        if msg['role'] == 'system':
            continue
        cell_width = 190 if len(msg['content']) > 100 else 200
        pdf.set_font("Arial", size=12)

        if msg['role'] == 'user':
            pdf.set_font("Arial", style="B", size=15)
        elif msg['role'] == 'assistant':
            pdf.set_font("Arial", style="", size=14)

        pdf.multi_cell(cell_width, 10, txt=f"{msg['role'].capitalize()}: {msg['content']}", align='L')

    pdf.output(pdf_filename)

def generate_docx(messages, docx_filename):
    document = Document()

    for msg in messages[1:]:
        if msg['role'] == 'system':
            continue
        sanitized_content = ''.join(char if 31 < ord(char) < 127 else ' ' for char in msg['content'])

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{msg['role'].capitalize()}: {sanitized_content}")

        if msg['role'] == 'assistant':
            run.font.size = Pt(12)
        elif msg['role'] == 'user':
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.size = Pt(14)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save(docx_filename)
